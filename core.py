#!/usr/bin/env python3
"""
AI Resume Tailoring - Core Module (Multi-User)
==============================================
业务逻辑层：多用户隔离、数据库管理、AI 客户端、Typst 编译、简历导入。
"""

import json
import os
import subprocess
import sys
import textwrap
from pathlib import Path
from datetime import datetime

import yaml
from openai import OpenAI

# ─── 路径常量 ───────────────────────────────────────────────────

def _is_frozen() -> bool:
    return getattr(sys, 'frozen', False)


def _get_resource_dir() -> Path:
    """只读资源目录（模板、字体、提示词、.env.yaml.example）。
    开发模式指向项目根目录，PyInstaller 打包后指向临时解压目录。"""
    if _is_frozen():
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


def _get_user_data_dir() -> Path:
    """可写用户数据目录（配置、数据库、输出文件）。
    开发模式使用项目目录下的 data/，打包后使用系统应用数据目录。"""
    if _is_frozen():
        if sys.platform == 'darwin':
            base = Path.home() / 'Library' / 'Application Support' / 'CV-Assistant'
        elif sys.platform == 'win32':
            base = Path(os.environ.get('APPDATA', str(Path.home()))) / 'CV-Assistant'
        else:
            base = Path.home() / '.local' / 'share' / 'CV-Assistant'
    else:
        base = Path(__file__).resolve().parent
    base.mkdir(parents=True, exist_ok=True)
    return base


def _get_typst_cmd() -> str:
    """返回 typst 命令。打包后使用捆绑的二进制，开发模式使用系统 PATH 中的。"""
    if _is_frozen():
        exe = 'typst.exe' if sys.platform == 'win32' else 'typst'
        # 按优先级查找：_MEIPASS → 可执行文件同目录 → PATH
        candidates = [Path(sys._MEIPASS) / exe]
        if sys.platform == 'darwin':
            candidates.append(Path(sys.executable).parent / exe)
        for cand in candidates:
            if cand.exists():
                # 确保可执行权限
                try:
                    os.chmod(cand, 0o755)
                except OSError:
                    pass
                return str(cand)
    return 'typst'


RESOURCE_DIR = _get_resource_dir()
USER_DATA_DIR = _get_user_data_dir()
DATA_DIR = USER_DATA_DIR / "data"
DEFAULT_PROMPT = RESOURCE_DIR / "prompts" / "system_prompt.md"

DEFAULT_MODEL = os.environ.get("CV_MODEL", "deepseek-chat")
DEFAULT_API_KEY = os.environ.get("CV_API_KEY", "")
DEFAULT_BASE_URL = os.environ.get("CV_BASE_URL", "https://api.deepseek.com")


# ═══════════════════════════════════════════════════════════════════
# 多用户路径
# ═══════════════════════════════════════════════════════════════════

def get_user_dir(user: str) -> Path:
    """返回用户的独立数据目录。"""
    return DATA_DIR / user


def get_user_db(user: str) -> Path:
    return get_user_dir(user) / "Master_Resume.yaml"


def get_user_outputs(user: str) -> Path:
    return get_user_dir(user) / "outputs"


def get_user_chat(user: str) -> Path:
    return get_user_dir(user) / "chat_history.json"


def get_user_jds(user: str) -> Path:
    return get_user_dir(user) / "jds"


def list_users() -> list:
    """列出所有已创建的用户。"""
    if not DATA_DIR.exists():
        return []
    users = []
    for d in sorted(DATA_DIR.iterdir()):
        if d.is_dir() and not d.name.startswith("."):
            db = d / "Master_Resume.yaml"
            users.append({
                "name": d.name,
                "has_db": db.exists(),
                "created": datetime.fromtimestamp(d.stat().st_mtime).strftime("%Y-%m-%d"),
            })
    return users


def create_user(name: str) -> Path:
    """创建新用户的数据目录和空模板。"""
    user_dir = get_user_dir(name)
    user_dir.mkdir(parents=True, exist_ok=True)
    (user_dir / "outputs").mkdir(exist_ok=True)
    (user_dir / "jds").mkdir(exist_ok=True)
    # 创建空数据库模板
    db_path = user_dir / "Master_Resume.yaml"
    if not db_path.exists():
        db_path.write_text(ResumeDatabase.init_template(), encoding="utf-8")
    # 创建空聊天记录
    chat_path = user_dir / "chat_history.json"
    if not chat_path.exists():
        chat_path.write_text("[]", encoding="utf-8")
    return user_dir


def delete_user(name: str) -> bool:
    """删除用户及所有数据。"""
    import shutil
    user_dir = get_user_dir(name)
    if user_dir.exists():
        shutil.rmtree(user_dir)
        return True
    return False


# ═══════════════════════════════════════════════════════════════════
# 全局配置
# ═══════════════════════════════════════════════════════════════════

def load_config() -> dict:
    config = {
        "api_key": DEFAULT_API_KEY,
        "base_url": DEFAULT_BASE_URL,
        "model": DEFAULT_MODEL,
    }
    config_file = USER_DATA_DIR / ".env.yaml"
    if config_file.exists():
        with open(config_file) as f:
            env_data = yaml.safe_load(f) or {}
            config.update(env_data)
    return config


def save_config(api_key: str, base_url: str, model: str) -> None:
    config_file = USER_DATA_DIR / ".env.yaml"
    data = {"api_key": api_key, "base_url": base_url, "model": model}
    with open(config_file, "w") as f:
        yaml.dump(data, f, allow_unicode=True)


def require_api_key(config: dict):
    if not config.get("api_key"):
        print("❌ 请设置 API Key：export CV_API_KEY='your-key'")
        sys.exit(1)


# ═══════════════════════════════════════════════════════════════════
# 数据库
# ═══════════════════════════════════════════════════════════════════

class ResumeDatabase:
    def __init__(self, user: str = "default"):
        self.user = user
        self.path = get_user_db(user)
        self.data = {}

    def exists(self) -> bool:
        return self.path.exists()

    def load(self) -> dict:
        if not self.path.exists():
            raise FileNotFoundError(f"数据库不存在: {self.path}")
        with open(self.path, encoding="utf-8") as f:
            self.data = yaml.safe_load(f) or {}
        return self.data

    def save(self, data: dict = None):
        if data is not None:
            self.data = data
        self.path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.path, "w", encoding="utf-8") as f:
            yaml.dump(self.data, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

    def to_yaml_string(self) -> str:
        return yaml.dump(self.data, allow_unicode=True, default_flow_style=False, sort_keys=False)

    def summary(self) -> dict:
        d = self.data
        return {
            "name": d.get("personal", {}).get("name", "未填写"),
            "email": d.get("personal", {}).get("email", ""),
            "education_count": len(d.get("education", [])),
            "work_count": len(d.get("work_experience", [])),
            "project_count": len(d.get("projects", [])),
            "skills_count": len(d.get("skills", {}).get("programming_languages", [])),
            "exists": True,
        }

    @staticmethod
    def init_template() -> str:
        return textwrap.dedent("""\
        # ============================================
        # 个人简历数据库
        # ============================================
        personal:
          name: ""
          email: ""
          phone: ""
          location: ""
          linkedin: ""
          github: ""
          website: ""

        education:
          - school: ""
            degree: ""
            major: ""
            date: ""
            gpa: ""
            highlights: []

        work_experience:
          - company: ""
            role: ""
            date: ""
            location: ""
            highlights: []

        projects:
          - name: ""
            role: ""
            date: ""
            tech_stack: ""
            url: ""
            description: ""
            highlights: []

        skills:
          programming_languages: []
          frameworks_and_tools: []
          domains: []
          languages:
            - "中文（母语）"

        publications: []
        certifications: []
        """)


# ═══════════════════════════════════════════════════════════════════
# AI 客户端
# ═══════════════════════════════════════════════════════════════════

class AIClient:
    def __init__(self, config: dict):
        base_url = config["base_url"].rstrip("/")
        if not base_url.endswith("/v1"):
            base_url += "/v1"
        self.client = OpenAI(api_key=config["api_key"], base_url=base_url)
        self.model = config["model"]

    def chat(self, system_prompt: str, user_message: str, temperature: float = 0.3, max_retries: int = 2) -> str:
        last_error = None
        for attempt in range(max_retries + 1):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    temperature=temperature,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message},
                    ],
                    timeout=120,
                )
                msg = response.choices[0].message
                return (msg.content or getattr(msg, "reasoning_content", "") or "").strip()
            except Exception as e:
                last_error = e
                if attempt < max_retries:
                    import time
                    time.sleep(2)
        raise last_error

    def chat_long(self, system_prompt: str, user_message: str, temperature: float = 0.3) -> str:
        """流式请求，避免长回复时的代理超时（如 interview prep / tailor）。"""
        try:
            full = []
            for token in self.chat_long_stream(system_prompt, user_message, temperature):
                full.append(token)
            return "".join(full).strip()
        except Exception:
            return self.chat(system_prompt, user_message, temperature)

    def chat_long_stream(self, system_prompt: str, user_message: str, temperature: float = 0.3):
        """流式请求的 generator 版本，用于 SSE 端点逐 token 推送。"""
        stream = self.client.chat.completions.create(
            model=self.model,
            temperature=temperature,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            stream=True,
            timeout=300,
        )
        for chunk in stream:
            if chunk.choices:
                content = getattr(chunk.choices[0].delta, "content", None)
                if content:
                    yield content

    def chat_stream(self, system_prompt: str, messages: list, temperature: float = 0.7):
        full_messages = [{"role": "system", "content": system_prompt}] + messages
        try:
            stream = self.client.chat.completions.create(
                model=self.model,
                temperature=temperature,
                messages=full_messages,
                stream=True,
            )
            for chunk in stream:
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta
                # 跳过 reasoning_content（推理模型的思考过程），只输出最终回复
                if getattr(delta, "content", None):
                    yield delta.content
        except Exception:
            response = self.client.chat.completions.create(
                model=self.model,
                temperature=temperature,
                messages=full_messages,
                stream=False,
            )
            yield response.choices[0].message.content or ""

    def tailor_pipeline(self, jd_text: str, db_yaml: str, prompt_template: str) -> str:
        system_prompt = prompt_template.replace("{jd_text}", jd_text).replace("{master_yaml}", db_yaml)
        result = self.chat_long(system_prompt, "请根据以上 JD 和简历数据库，生成定制化的 Typst 简历代码。")
        return _ensure_typst_import(result)

    def interview_prep_pipeline(self, jd_text: str, db_yaml: str) -> dict:
        system_prompt = INTERVIEW_PREP_PROMPT.replace("{jd_text}", jd_text).replace("{master_yaml}", db_yaml)
        result = self.chat_long(system_prompt, "请根据以上 JD 和简历数据库，输出面试准备分析报告（JSON 格式）。")
        # 清理可能的 markdown fence
        result = _clean_code_fence(result, "json")
        try:
            return json.loads(result)
        except json.JSONDecodeError:
            return {"raw": result, "parse_error": True}

    def test_connection(self) -> tuple[bool, str]:
        try:
            self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": "回复 OK"}],
                max_tokens=10,
            )
            return True, "连接成功"
        except Exception as e:
            return False, str(e)


# ═══════════════════════════════════════════════════════════════════
# Typst 编译
# ═══════════════════════════════════════════════════════════════════

class TypstCompiler:
    def __init__(self, user: str = "default"):
        self.output_dir = get_user_outputs(user)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def compile(self, typst_code: str, output_name: str = "tailored_cv") -> tuple[Path, bool, str]:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = f"{output_name}_{ts}"
        typst_path = self.output_dir / f"{safe_name}.typ"
        pdf_path = self.output_dir / f"{safe_name}.pdf"

        typst_path.write_text(typst_code, encoding="utf-8")

        typst_cmd = _get_typst_cmd()
        _fix_bundled_binary(typst_cmd)
        if not _check_command(typst_cmd):
            return pdf_path, False, "Typst CLI 未安装。请在左侧系统面板点击「一键安装 Typst」，或手动运行: brew install typst"

        fonts_dir = RESOURCE_DIR / "fonts"
        cmd = [typst_cmd, "compile"]
        if fonts_dir.is_dir():
            cmd += ["--font-path", str(fonts_dir)]
        cmd += [str(typst_path), str(pdf_path)]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            return pdf_path, True, str(pdf_path)
        else:
            return pdf_path, False, result.stderr

    def typst_available(self) -> bool:
        return _check_command("typst")


def _check_command(cmd: str) -> bool:
    try:
        subprocess.run([cmd, "--version"], capture_output=True, check=True)
        return True
    except Exception:
        return False


def check_typst_available() -> dict:
    """检查 Typst 是否可用，返回详细信息。"""
    cmd = _get_typst_cmd()
    _fix_bundled_binary(cmd)
    available = _check_command(cmd)
    result = {
        "available": available,
        "path": cmd,
        "is_bundled": _is_frozen() and cmd != "typst",
    }
    if available:
        try:
            r = subprocess.run([cmd, "--version"], capture_output=True, text=True)
            result["version"] = r.stdout.strip()
        except Exception:
            pass
    return result


def install_typst() -> tuple[bool, str]:
    """一键安装 Typst CLI。macOS 用 brew，否则引导用户手动安装。"""
    if sys.platform == "darwin":
        # 优先尝试 brew install
        for installer in ["/opt/homebrew/bin/brew", "/usr/local/bin/brew", "brew"]:
            try:
                r = subprocess.run(
                    [installer, "install", "typst"],
                    capture_output=True, text=True, timeout=300,
                )
                if r.returncode == 0:
                    return True, "Typst 已通过 Homebrew 安装成功"
            except Exception:
                continue
        return False, (
            "自动安装失败。请手动安装 Typst：\n"
            "1. 打开终端运行: brew install typst\n"
            "2. 或从 https://github.com/typst/typst/releases 下载"
        )
    elif sys.platform == "win32":
        return False, (
            "请手动安装 Typst：\n"
            "1. 打开 https://github.com/typst/typst/releases\n"
            "2. 下载 typst-x86_64-pc-windows-msvc.zip\n"
            "3. 解压后将 typst.exe 放到系统 PATH 中"
        )
    else:
        return False, "请从 https://github.com/typst/typst/releases 手动安装 Typst"


def _fix_bundled_binary(cmd: str) -> bool:
    """确保捆绑的二进制可执行。macOS 下载的 DMG 可能有 quarantine 标记。"""
    if not _is_frozen() or sys.platform != 'darwin':
        return True
    p = Path(cmd)
    try:
        # 确保有可执行权限
        os.chmod(p, 0o755)
        # 尝试移除 quarantine 属性（忽略无 xattr 命令的错误）
        subprocess.run(
            ['xattr', '-d', 'com.apple.quarantine', str(p)],
            capture_output=True,
        )
    except Exception:
        pass
    return True


# ═══════════════════════════════════════════════════════════════════
# 简历导入
# ═══════════════════════════════════════════════════════════════════

def import_resume(config: dict, file_path: str, user: str = "default") -> dict:
    fp = Path(file_path)
    if not fp.exists():
        return {"success": False, "message": f"文件不存在: {file_path}"}

    ext = fp.suffix.lower()
    text = ""

    try:
        if ext == ".txt":
            text = fp.read_text(encoding="utf-8")
        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(fp))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(str(fp))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            return {"success": False, "message": f"不支持的文件格式: {ext}（支持: .txt, .pdf, .docx）"}
    except ImportError as e:
        return {"success": False, "message": f"缺少依赖库: {e}"}

    if not text.strip():
        return {"success": False, "message": "未能从文件中提取到文本内容。"}

    # 预处理：清洗 PDF 提取的文本，整理成结构化 Markdown
    cleaned_text = _preprocess_resume_text(text, ext)

    parse_prompt = textwrap.dedent("""\
    你是一个简历解析专家。请将以下简历文本解析为 YAML 格式。
    提取以下字段（尽可能完整，没有的信息写空或省略）：

    personal:
      name: 姓名 / email: 邮箱 / phone: 电话 / location: 城市
    education:
      - school: 学校 / degree: 学位 / major: 专业 / date: 时间 / highlights: [亮点]
    work_experience:
      - company: 公司 / role: 职位 / date: 时间 / location: 地点 / highlights: [描述, ...]
    projects:
      - name: 项目名 / role: 角色 / date: 时间 / tech_stack: 技术栈 / description: 简介 / highlights: [描述, ...]
    skills:
      programming_languages: [] / frameworks_and_tools: [] / domains: [] / languages: []

    要求：
    1. 输出纯 YAML，不要用 markdown 代码块包裹
    2. highlights 中的每条描述保留原意，以动词开头
    3. 不要编造任何原文没有的信息
    """)

    client = AIClient(config)
    result = client.chat(parse_prompt, f"<RESUME_TEXT>\n{cleaned_text}\n</RESUME_TEXT>")
    yaml_text = _clean_code_fence(result, "yaml")

    try:
        data = yaml.safe_load(yaml_text)
        if not isinstance(data, dict):
            return {"success": False, "message": f"AI 输出格式异常（非 YAML 字典）", "raw": result}

        db = ResumeDatabase(user)
        if db.exists():
            existing = yaml.safe_load(db.path.read_text(encoding="utf-8")) or {}
            if not isinstance(existing, dict):
                existing = {}
            merged = _merge_data(existing, data)
            db.save(merged)
        else:
            db.save(data)
        return {"success": True, "message": "简历内容已解析并导入", "data": (db.data if db.exists() else data)}
    except yaml.YAMLError as e:
        return {"success": False, "message": f"AI 输出的 YAML 解析失败: {e}", "raw": result}
    except Exception as e:
        return {"success": False, "message": f"合并数据出错: {e}", "raw": result}


def import_resume_stream(config: dict, file_path: str, user: str = "default"):
    """流式版本的简历导入，通过 SSE 推送进度，避免长请求超时。"""
    import json as _json
    import sys as _sys

    def _log(msg):
        try:
            with open(Path.home() / ".cv-assistant-debug.log", "a") as f:
                f.write(f"[{datetime.now().isoformat()}] {msg}\n")
        except Exception:
            pass

    _log(f"import_resume_stream start: {file_path}")

    fp = Path(file_path)
    if not fp.exists():
        _log(f"file not found: {file_path}")
        yield f"data: {_json.dumps({'error': '文件不存在'})}\n\n"
        return

    ext = fp.suffix.lower()
    _log(f"file ext: {ext}, size: {fp.stat().st_size}")

    # 阶段1: 提取文本
    yield f"data: {_json.dumps({'stage': 'extract', 'message': '正在读取文件...'})}\n\n"
    text = ""
    try:
        if ext == ".txt":
            text = fp.read_text(encoding="utf-8")
        elif ext == ".pdf":
            _log("importing pypdf...")
            from pypdf import PdfReader
            _log("PdfReader imported, opening...")
            reader = PdfReader(str(fp))
            _log(f"PDF opened, {len(reader.pages)} pages")
            pages_text = []
            for i, page in enumerate(reader.pages):
                pages_text.append(page.extract_text() or "")
                if i % 5 == 0:
                    yield f"data: {_json.dumps({'stage': 'extract', 'message': f'正在解析 PDF 第 {i+1} 页...'})}\n\n"
            text = "\n".join(pages_text)
            _log(f"PDF text extracted, {len(text)} chars")
        elif ext in (".docx", ".doc"):
            _log("importing docx...")
            from docx import Document
            doc = Document(str(fp))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            yield f"data: {_json.dumps({'error': f'不支持的文件格式: {ext}'})}\n\n"
            return
    except ImportError as e:
        _log(f"ImportError: {e}")
        yield f"data: {_json.dumps({'error': f'缺少依赖库: {e}'})}\n\n"
        return
    except Exception as e:
        _log(f"text extraction error: {e}")
        import traceback as _tb
        yield f"data: {_json.dumps({'error': f'文件解析失败: {e}', 'traceback': _tb.format_exc()})}\n\n"
        return

    if not text.strip():
        yield f"data: {_json.dumps({'error': '未能从文件中提取到文本内容。'})}\n\n"
        return

    # 预处理
    yield f"data: {_json.dumps({'stage': 'preprocess', 'message': '正在清洗文本...'})}\n\n"
    cleaned_text = _preprocess_resume_text(text, ext)

    # 阶段2: AI 解析（流式）
    yield f"data: {_json.dumps({'stage': 'ai_start', 'message': 'AI 正在解析简历...'})}\n\n"

    parse_prompt = textwrap.dedent("""\
    你是一个简历解析专家。请将以下简历文本解析为 YAML 格式。
    提取以下字段（尽可能完整，没有的信息写空或省略）：

    personal:
      name: 姓名 / email: 邮箱 / phone: 电话 / location: 城市
    education:
      - school: 学校 / degree: 学位 / major: 专业 / date: 时间 / highlights: [亮点]
    work_experience:
      - company: 公司 / role: 职位 / date: 时间 / location: 地点 / highlights: [描述, ...]
    projects:
      - name: 项目名 / role: 角色 / date: 时间 / tech_stack: 技术栈 / description: 简介 / highlights: [描述, ...]
    skills:
      programming_languages: [] / frameworks_and_tools: [] / domains: [] / languages: []

    要求：
    1. 输出纯 YAML，不要用 markdown 代码块包裹
    2. highlights 中的每条描述保留原意，以动词开头
    3. 不要编造任何原文没有的信息
    """)

    client = AIClient(config)
    full_result = ""
    try:
        for token in client.chat_long_stream(parse_prompt, f"<RESUME_TEXT>\n{cleaned_text}\n</RESUME_TEXT>"):
            full_result += token
            yield f"data: {_json.dumps({'stage': 'ai_token', 'token': token})}\n\n"
    except Exception as e:
        _log(f"AI call error: {e}")
        import traceback as _tb
        yield f"data: {_json.dumps({'error': f'AI 调用失败: {e}', 'traceback': _tb.format_exc()})}\n\n"
        return
    except Exception as e:
        yield f"data: {_json.dumps({'error': f'AI 调用失败: {e}'})}\n\n"
        return

    # 阶段3: 解析 & 保存
    yield f"data: {_json.dumps({'stage': 'saving', 'message': '正在解析并保存...'})}\n\n"
    yaml_text = _clean_code_fence(full_result, "yaml")

    try:
        data = yaml.safe_load(yaml_text)
        if not isinstance(data, dict):
            yield f"data: {_json.dumps({'error': 'AI 输出格式异常', 'raw': full_result})}\n\n"
            return

        db = ResumeDatabase(user)
        if db.exists():
            existing = yaml.safe_load(db.path.read_text(encoding="utf-8")) or {}
            if not isinstance(existing, dict):
                existing = {}
            merged = _merge_data(existing, data)
            db.save(merged)
        else:
            db.save(data)

        summary = db.summary() if db.exists() else {"name": data.get("personal", {}).get("name", "")}
        yield f"data: {_json.dumps({'stage': 'done', 'message': '简历解析完成！', 'summary': summary})}\n\n"
    except yaml.YAMLError as e:
        _log(f"YAML error: {e}")
        yield f"data: {_json.dumps({'error': f'YAML 解析失败: {e}', 'raw': full_result[:500]})}\n\n"
    except Exception as e:
        _log(f"save error: {e}")
        import traceback as _tb
        yield f"data: {_json.dumps({'error': f'保存出错: {e}', 'traceback': _tb.format_exc(), 'raw': full_result[:500]})}\n\n"
    _log("import_resume_stream done")


# ═══════════════════════════════════════════════════════════════════
# 对话式录入
# ═══════════════════════════════════════════════════════════════════

CHAT_SYSTEM_PROMPT = textwrap.dedent("""\
你是一名资深职业规划师（GCDF）兼简历顾问。你的任务不是机械地收集信息，而是像一位专业的 career coach 一样，通过深度对话帮助用户发现、梳理并提炼他们的职业经历，最终形成一份结构化的简历数据库。

## 咨询理念

- **以来询者为中心**：你不是在审问，而是在陪伴用户共同探索。你的角色是镜子，帮用户看到自己可能忽略的亮点。
- **给予多于索求**：每次交流中，你不仅要提问，更要提供洞察——帮用户提炼、归纳、升华他们刚才说的内容。
- **非评判性态度**：不对用户的经历做价值判断，而是帮他们发现每段经历中的独特价值。
- **助人自助**：你的最终目标是帮用户建立起对自己职业竞争力的清晰认知，而不仅是填完一份表格。

## 对话四阶段

### 阶段一：建立关系 & 明确目标（收纳面谈）
- 用开放式问题开场，了解用户当前求职方向或职业困惑
- 了解用户想通过这份简历达到什么目的（转行？晋升？入行？）
- 示例：「你好！很高兴能帮你梳理简历。在开始之前，我想先了解一下——你目前在看什么样的机会？或者说，你希望通过这份简历达成什么目标呢？」

### 阶段二：分模块深度挖掘（一次只深入一个领域）
按以下顺序逐一深入，每轮只聚焦一个模块：

1. **基本信息**：姓名、联系方式、所在地、LinkedIn/GitHub（如有）
2. **教育背景**：学校、学位、专业、起止时间、GPA/荣誉、核心课程
3. **工作经历**：公司、职位、起止时间、工作内容
4. **项目经历**：项目名称、你的角色、起止时间、技术栈、具体贡献
5. **技能图谱**：编程语言、框架工具、领域知识、语言能力
6. **学术成果**：论文发表、专利、证书、竞赛获奖

### 阶段三：成就萃取（对每条经历追问深度）
这是最关键的阶段。用户提到任何经历时，运用以下提问技巧挖掘深层内容：

**成果导向追问**：
- 「这件事的结果是什么？有没有可以量化的数据？」（推动量化）
- 「如果我是HR，我怎么知道你把这件事做得好？」（区分平庸与优秀）
- 「在这个过程中，你个人做了什么（而非团队做了什么）？」（区分 I vs We）

**困难与挑战追问**：
- 「做这件事的过程中，遇到过什么特别的困难或挑战吗？」（挖掘冲突感）
- 「当时资源/时间/信息是否充足？你是怎么应对的？」（展示应变能力）

**亮点发现追问**：
- 「这段经历中，最让你有成就感的是什么？」（发现热情所在）
- 「跟其他做类似事情的人相比，你觉得你做得不一样的地方在哪里？」（发现独特优势）
- 「做这件事的时候，你的感受是什么？」（挖掘动力与价值观）

**可迁移能力识别**：
- 「这件事锻炼了你哪些能力？这些能力对你目标岗位有什么帮助？」（连接经历与目标）
- 帮用户将技能归类为三类：专业知识技能（名词）、可迁移技能（动词）、自我管理技能（形容词）

### 阶段四：综合整理 & 输出
- 用户说「完成」时，汇总所有对话内容，输出完整的 YAML 格式简历数据库
- 输出前，可以给用户一个简要总结：「好的，根据我们的对话，我帮你梳理了以下内容……确认无误后我就生成 YAML 文件。」

## 对话技巧规则

1. **每次只问一个问题**，等用户回答后再继续
2. **先镜射再追问**：用户回答后，先用一两句话提炼/肯定（「明白了，你在这个项目中不仅负责了XX，还主动推动了YY，对吧？」），再自然过渡到下一个问题
3. **用开放式问题替代封闭式问题**：
   - ❌ 「你有量化数据吗？」
   - ✅ 「这件事的结果可以用什么数字来衡量吗？比如影响了多少人、提升了多少效率、节省了多少时间？」
4. **帮用户补全信息**：如果用户说了模糊的内容，温和地追问具体细节（「你提到『效率提升了』，具体是从多少提升到了多少？大概节省了多少时间？」）
5. **适时换位思考**：引导用户站在HR/用人部门角度审视自己的经历（「假设你是这个岗位的招聘经理，你看到这段描述会觉得这个候选人有什么独特价值？」）
6. **保持对话自然流畅**，不要像在填表格
7. **识别并记录可迁移能力**：即使用户的某段经历与目标岗位不直接相关，也要帮用户看到其中的可迁移价值

## YAML 输出结构

当用户说「完成」时，输出如下结构的 YAML：

```yaml
personal:
  name: ""
  email: ""
  phone: ""
  location: ""
  linkedin: ""
  github: ""
education:
  - school: ""
    degree: ""
    major: ""
    date: ""
    highlights: []
work_experience:
  - company: ""
    role: ""
    date: ""
    highlights: []
projects:
  - name: ""
    role: ""
    date: ""
    tech_stack: []
    highlights: []
skills:
  programming_languages: []
  frameworks_and_tools: []
  domains: []
  languages: []
publications:
  - title: ""
    journal: ""
    date: ""
    role: ""
certifications: []
```

## 绝对禁止（输出格式要求）

**你的每条回复必须是对用户说的最终对话内容，直接说人话。严禁输出以下内容：**
- 思考过程、推理过程、分析过程
- "首先…"/"等下…"/"不对…"/"哦不对…"/"我再想想…"这类自言自语
- 任何括号或引号内的内心独白
- 对你行为的描述（如"我来总结一下"、"我先肯定一下"）
- Markdown 格式标记（除非在最终输出 YAML 时）

请用中文对话，像一位专业、温和、善于倾听的职场前辈。""")


INTERVIEW_PREP_PROMPT = textwrap.dedent("""\
你是一位资深的面试辅导教练兼招聘专家。你的任务是基于目标 JD 和候选人的简历数据库，输出一份完整的面试准备报告。

## 输出格式

严格输出 JSON（不要 markdown fences），结构如下：

```json
{
  "gap_analysis": [
    {"requirement": "JD 要求的能力/经验", "match_level": "强匹配/部分匹配/缺失", "evidence": "候选人现有经历中的证据或空缺说明", "importance": "高/中/低"}
  ],
  "gap_remediation": [
    {"gap": "缺失的能力/经验", "how_to_fill": "如何快速弥补（学习资源、项目建议、话术应对）", "alternative_angle": "如何用现有经历侧面回应此要求"}
  ],
  "interview_questions": [
    {"category": "技术能力/行为面试/情景题/动机与文化", "question": "面试官可能问的问题", "why_ask": "面试官为什么问这个", "answer_tip": "回答要点提示（结合候选人真实经历）"}
  ],
  "self_intro": {
    "duration": "1-2 分钟版本",
    "script": "完整自我介绍文案，以第一人称撰写",
    "structure": "开场问候 → 背景与核心竞争力概述 → 2-3 个与 JD 最相关的成就亮点 → 求职动机 → 结束语"
  }
}
```

## 分析要求

1. **gap_analysis**：逐条拆解 JD 中的核心要求（硬技能、软技能、经验年限、行业背景、学历等），与简历数据库逐一对比，标注匹配程度。最多 10 条。
2. **gap_remediation**：针对「部分匹配」和「缺失」的能力，给出具体可行的弥补建议。每个建议都要包含：
   - 如何快速补充（课程/书籍/实战项目）
   - 面试中如何用现有经历"侧面回应"这个问题
3. **interview_questions**：模拟面试官视角，预测 10-15 个面试问题，覆盖技术、行为、情景、动机 4 个维度。每个问题都要说明考察点，并给出结合候选人真实经历的回答要点。
4. **self_intro**：生成一段 1-2 分钟（约 250-350 字）的中文自我介绍。要求：
   - 用第一人称
   - 开场简洁有力
   - 选取与 JD 最相关的 2-3 个成就亮点
   - 自然表达求职动机
   - 以开放式结束语收尾（引导面试官追问）
   - 所有引用的事例/数据必须来自简历数据库，不得虚构

## 核心原则

- **所有建议必须基于候选人的真实经历**，不得建议没有做过的事情
- **回答要点要具体**，给出话术方向的提示，而不是泛泛地说「准备一下」
- **自我介绍要自然口语化**，不要像在读简历
- **用中文输出**，但技术术语可用英文

<TARGET_JD>
{jd_text}
</TARGET_JD>

<MASTER_DATABASE>
{master_yaml}
</MASTER_DATABASE>""")


def chat_extract_yaml(config: dict, messages: list, user: str = "default") -> dict:
    client = AIClient(config)
    msgs = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}] + messages
    msgs.append({"role": "user", "content": "请根据我们刚才的对话，输出完整的 YAML 格式简历数据库。"})
    response = client.client.chat.completions.create(
        model=client.model,
        messages=msgs,
        temperature=0.3,
    )
    yaml_text = _clean_code_fence(response.choices[0].message.content, "yaml")
    try:
        # 使用 safe_load_all 处理 AI 可能输出的多文档 YAML（用 --- 分隔）
        docs = list(yaml.safe_load_all(yaml_text))
        data = {}
        for doc in docs:
            if isinstance(doc, dict):
                data = _merge_data(data, doc)
        db = ResumeDatabase(user)
        if db.exists():
            db.load()
            data = _merge_data(db.data, data)
        db.save(data)
        return {"success": True, "message": "简历数据库已保存", "summary": db.summary()}
    except yaml.YAMLError as e:
        return {"success": False, "message": f"YAML 解析失败: {e}", "raw": yaml_text}


# ═══════════════════════════════════════════════════════════════════
# 聊天记录持久化
# ═══════════════════════════════════════════════════════════════════

def load_chat_history(user: str) -> list:
    p = get_user_chat(user)
    if p.exists():
        with open(p, encoding="utf-8") as f:
            return json.load(f)
    return []


def save_chat_history(user: str, messages: list):
    p = get_user_chat(user)
    p.parent.mkdir(parents=True, exist_ok=True)
    with open(p, "w", encoding="utf-8") as f:
        json.dump(messages, f, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════════════════
# JD 管理
# ═══════════════════════════════════════════════════════════════════

def save_jd(user: str, name: str, content: str):
    jd_dir = get_user_jds(user)
    jd_dir.mkdir(parents=True, exist_ok=True)
    (jd_dir / f"{name}.txt").write_text(content, encoding="utf-8")


def list_jds(user: str) -> list:
    jd_dir = get_user_jds(user)
    if not jd_dir.exists():
        return []
    jds = []
    for f in sorted(jd_dir.glob("*.txt")):
        jds.append({
            "name": f.stem,
            "preview": f.read_text(encoding="utf-8")[:200],
            "time": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
        })
    return jds


def get_jd(user: str, name: str) -> str:
    jd_path = get_user_jds(user) / f"{name}.txt"
    if jd_path.exists():
        return jd_path.read_text(encoding="utf-8")
    return ""


# ═══════════════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════════════

def _preprocess_resume_text(text: str, ext: str) -> str:
    """清洗从 PDF/DOCX 提取的原始文本，整理成结构化 Markdown。

    PDF 提取的文本常有：多余空行、断行错位、页码混入、空格异常等问题。
    预处理后发给 AI 解析更快更准。
    """
    import re

    # ── 通用清洗 ──
    # 统一换行符
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    if ext == ".pdf":
        # PDF 提取的文本几乎每行都是独立段落，需要合并被切断的句子
        lines = text.split("\n")
        cleaned = []
        buf = ""
        for line in lines:
            s = line.strip()
            if not s:
                if buf:
                    cleaned.append(buf)
                    buf = ""
                cleaned.append("")
                continue
            # 全是非字母数字符号的行 → 可能是装饰线，跳过
            if len(re.sub(r"[^\w一-鿿]", "", s)) < 3:
                continue
            # 页码行：纯数字或 "第X页"
            if re.match(r"^(第?\s*\d+\s*页?|\d+/\d+)$", s):
                continue
            # 上一行以标点结尾 → 独立段落
            if buf and re.search(r"[。！？.!?，,；;：:—\-—…]$", buf):
                cleaned.append(buf)
                buf = s
            # 新行以大写/中文开头且 buf 有意义 → 可能新段落
            elif buf and (s[0].isupper() or "一" <= s[0] <= "鿿" or s[0] == "●" or s[0] == "-"):
                cleaned.append(buf)
                buf = s
            else:
                buf = (buf + " " + s) if buf else s
        if buf:
            cleaned.append(buf)
        text = "\n".join(cleaned)

    # ── 通用规范化 ──
    # 压缩多个空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去掉每行首尾多余空格
    lines = [l.strip() for l in text.split("\n")]
    text = "\n".join(lines)
    # 压缩多余空格
    text = re.sub(r" {2,}", " ", text)

    return text.strip()


def _clean_code_fence(text: str, lang: str = "") -> str:
    text = text.strip()
    if lang and text.startswith(f"```{lang}"):
        text = text[len(f"```{lang}"):]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


def _ensure_typst_import(code: str) -> str:
    """确保 Typst 代码包含 brilliant-cv import 语句。"""
    import re
    code = code.strip()
    if not re.search(r'#import\s+"@preview/brilliant-cv', code):
        code = '#import "@preview/brilliant-cv:2.0.0": *\n\n' + code
    return code


def _patch_typst_metadata(code: str) -> str:
    """补全 brilliant-cv v2.0.0 metadata 中缺失的必要键。

    这些键在 cvSection/cvEntry 内部通过 metadata.layout.at() 访问，
    如果键不存在，default 值是长度类型而非字符串，eval() 会报错。
    """
    import re

    # 确保 layout 包含 before_section_skip 等必要键
    layout_fixes = [
        ('before_section_skip: ', 'before_section_skip: "1pt",'),
        ('before_entry_skip: ', 'before_entry_skip: "1pt",'),
        ('before_entry_description_skip: ', 'before_entry_description_skip: "1pt",'),
    ]

    for key_pattern, replacement in layout_fixes:
        if key_pattern.rstrip(": ") not in code:
            # 在 header: 之前插入缺失的键
            code = re.sub(
                r'(layout:\s*\(\s*)',
                r'\1' + replacement + '\n    ',
                code,
                count=1,
            )

    # 确保 inject 存在
    if 'inject:' not in code:
        code = re.sub(
            r'(lang:\s*\()',
            'inject: (\n    inject_ai_prompt: false,\n    inject_keywords: false,\n    injected_keywords_list: (),\n  ),\n  \1',
            code,
            count=1,
        )

    # 确保 non_latin 存在
    if 'non_latin:' not in code:
        code = re.sub(
            r'(zh:\s*\([^)]*cv_footer[^)]*\),)',
            r'\1\n    non_latin: (name: "", font: "Heiti SC"),',
            code,
        )

    # 转义内容块中的 #：C#、F# 等语言名中的 # 会被 Typst 误解析为代码表达式起始符。
    # 只转义后面不跟字母/下划线/左括号的 #（保留 #hBar() #metadata.field 等合法调用）
    code = re.sub(r'([A-Za-z])#(?![a-zA-Z_(])', r'\1\\#', code)

    # AI 偶尔会输出全角标点作为代码分隔符（如中文括号 ），修复为半角
    code = code.replace('）', ')')
    code = code.replace('（', '(')

    return code


def _merge_data(base: dict, new: dict) -> dict:
    for key, val in new.items():
        if key not in base or base[key] is None:
            base[key] = val
        elif isinstance(val, dict):
            if base[key] is None:
                base[key] = {}
            for sk, sv in val.items():
                if sv:
                    base[key][sk] = sv
        elif isinstance(val, list):
            if base[key] is None:
                base[key] = []
            for item in val:
                if item not in base[key]:
                    base[key].append(item)
    return base


def list_outputs(user: str = "default") -> list:
    out = get_user_outputs(user)
    if not out.exists():
        return []
    pdfs = sorted(out.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
    return [{"name": p.name, "path": str(p), "time": datetime.fromtimestamp(p.stat().st_mtime).strftime("%Y-%m-%d %H:%M")} for p in pdfs]
